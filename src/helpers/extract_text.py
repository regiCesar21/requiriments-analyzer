import fitz  # PyMuPDF
import pandas as pd
import re
from os.path import join
from collections import defaultdict
from concurrent.futures import ProcessPoolExecutor
import multiprocessing

import nltk
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer, util
import torch

nltk.download('punkt')

# Compila os padrões uma vez para reuso (melhora o desempenho se for usado em loop)
_clean_patterns = [
    (re.compile(r'http[s]?://\S+'), ''),  # URLs
    (re.compile(r'\((?:[^\(\)]*?\d{4}[^\(\)]*?;?\s*)+\)'), ''),  # Citações (Silva, 2003)
    (re.compile(r'©\s?\d{4}\s?The Authors', re.IGNORECASE), ''),  # Copyright
    (re.compile(r'\[\d+\]'), ''),  # Citações [1], [12]
    (re.compile(r'\s+([?.!,;:])'), r'\1'),  # Espaço antes de pontuação
    (re.compile(r'\s+'), ' '),  # Espaços múltiplos
]

def clean_text(text):
    if not text:
        return ""
    for pattern, repl in _clean_patterns:
        text = pattern.sub(repl, text)
    return text.strip()

def clean_texts_parallel(text, max_workers=None):
    print(f"| ### 🧹 Limpando texto extraído... ### |")
    text_list = split_text_by_words(text)
    if max_workers is None:
        max_workers = multiprocessing.cpu_count()
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        cleaned = list(executor.map(clean_text, text_list))
    print(f"| ### ✨ Text ready/finalized after cleaning... ### |")
    return cleaned

def clear_texts(text):
    # print()
    # print('| ----------------------------------------------- TEXTOS ANTES DA LIMPEZA -------------------------------------------------------- |')
    # print(text)
    # print('| -------------------------------------------------------------------------------------------------------------------------------- |')

    if pd.isnull(text):
        return None

    # Remover espaços extras
    text = text.strip()

    # Remover apenas números ou palavras numéricas
    if re.fullmatch(r"\(?\d{1,4}\)?", text):  # Ex: "1980", "(1980)", "123"
        return None

    # Remover textos curtos sem sentido (1-2 palavras)
    if len(text.split()) < 3:
        return None

    # Remover padrões tipo: Journal info, volume, páginas — ex: "Information Systems Research, 21, 748-59"
    if re.search(r"\b\d{1,2},\s*\d{2,3}(-\d{1,3})?\b", text):
        return None

    # Remover frases genéricas como "To estimate", "To assess"
    if re.fullmatch(r"To\s+\w+", text, re.IGNORECASE):
        return None

    return text

# 1. Extrair texto do PDF
def extract_pdf_text(path_pdf):
    print(f"| ### 📄 Starting PDF text extraction... ### |")
    try:
        doc = fitz.open(path_pdf)
        meta_doc = doc.metadata
        metadata_str = format_metadata(meta_doc)
        text_total = ""
        for page in doc:
            blocks = page.get_text("blocks")
            for block in blocks:
                text_total += clear_texts(block[4])
        doc.close()
        print(f"| ### 📑 Finished PDF text extraction... ### |")
        return text_total, metadata_str
    except Exception as e:
        print(f"An error occurred while extracting the text: {e}")
        return None

def extrair_texto_txt(caminho_txt):
    try:
        with open(caminho_txt, 'r', encoding='utf-8') as arquivo:
            texto_total = arquivo.read()
        return texto_total
    except Exception as e:
        print(f"Ocorreu um erro ao ler o arquivo TXT: {e}")
        return None

# 2. Filtrar frases com base em palavras-chave
def extrair_frases_relevantes(texto, palavras_chave):
    frases = re.split(r'[.\n]', texto)
    relevantes = [
        frase.strip() for frase in frases
        if any(palavra.lower() in frase.lower() for palavra in palavras_chave) and len(frase.strip()) > 40
    ]
    return relevantes


def extract_relevant_phrases2(text, key_words, min_len=30, return_keywords=False):
    phrases = sent_tokenize(text)
    relevants = []

    key_words = list(set([p.lower() for p in key_words]))
    for phrase in phrases:
        phrase_stripped = phrase.strip()
        phrase_lower = phrase_stripped.lower()

        # Ignorar frases curtas
        if len(phrase_stripped) < min_len:
            continue
        # Ignorar frases com links ou padrões comuns de citação
        if re.search(r'https?://|www\.|doi\.org|et al\.|\(\d{4}\)|\[\d+\]', phrase_lower):
            continue
        # Ignorar frases que terminam com vírgula ou sem pontuação
        if not re.search(r'[.!?]$', phrase_stripped):
            continue
        # Verificar presença de palavras-chave
        matched_keywords = [word for word in key_words if re.search(r'\b' + re.escape(word) + r'\b', phrase_lower)]
        if matched_keywords:
            if return_keywords:
                relevants.append((phrase_stripped, matched_keywords))
            else:
                relevants.append(phrase_stripped)
    return relevants

def salvar_txt(texto):
    try:
        with open("texto.txt", 'w', encoding='utf-8') as f:
            f.write(texto)
        print("Arquivo 'texto.txt' salvo com sucesso!")
    except Exception as e:
        print(f"Ocorreu um erro ao salvar o arquivo: {e}")

def extract_key_words(kw_model, text):
    print(f"| ### 🔍 Starting keyword extraction... ### |")
    parts = split_text_by_words(text, size=500)
    all_keywords = []
    grouped_keywords = defaultdict(list)
    for i, part in enumerate(parts):
        print(f"\n🔹 Part {i + 1}")
        keywords = kw_model.extract_keywords(
            part,
            keyphrase_ngram_range=(1, 3),
            stop_words='english',
            top_n=5
        )
        for word, score in keywords:
            print(f"  - {word} (score: {round(score, 4)})")
            grouped_keywords[word].append(score)

    # === Sum the scores per keyword ===
    final_keywords = []
    for word, scores in grouped_keywords.items():
        total_score = sum(scores)
        final_keywords.append((word, round(total_score, 4)))

    # === Sort and display top keywords ===
    top_keywords = sorted(final_keywords, key=lambda x: x[1], reverse=True)
    print(f"| ### 🔑 Keywords extracted... ### |")
    return top_keywords[:20]

# === Split into parts of 500 words ===
def split_text_by_words(text, size=500):
    words = text.split()
    return [" ".join(words[i:i + size]) for i in range(0, len(words), size)]

def extract_relevant_phrases(text, only_keywords, model):
    print(f"| ### 🧠 Extracting relevant phrases based on keywords... ### |")
    phrases = sent_tokenize(text)
    # only_keywords = [kw[0] for kw in key_words]

    # model = SentenceTransformer('all-MiniLM-L6-v2')
    emb_phrases = model.encode(phrases, convert_to_tensor=True)
    emb_keywords = model.encode(only_keywords, convert_to_tensor=True)
    scores = util.cos_sim(emb_phrases, emb_keywords).mean(dim=1)
    k = min(400, scores.shape[0])
    top_idxs = torch.argsort(scores, descending=True)

    all_phrases = []
    for idx in top_idxs:
        print(f"- {phrases[idx]}")
        all_phrases.append(phrases[idx])

    print(f"| ### 📌 Relevant phrases extracted... ### |")
    return all_phrases

def format_metadata(meta: dict) -> str:
    fields = ["title", "author", "creationDate", "modDate", "subject", "keywords"]
    return "; ".join(f"{field}: {meta[field]}" for field in fields if field in meta and meta[field])


def _extract_docx(raw_bytes: bytes) -> str:
    """Extrai texto de um arquivo DOCX (bytes). Requer python-docx."""
    try:
        import docx as python_docx
    except ImportError:
        raise ValueError(
            "Suporte a DOCX requer a biblioteca 'python-docx'.\n"
            "Instale com: pip install python-docx"
        )
    from io import BytesIO
    doc = python_docx.Document(BytesIO(raw_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Inclui texto de tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_from_upload(uploaded_file) -> tuple[str, list[str]]:
    """Extrai texto e sentenças de um UploadedFile do Streamlit (PDF, TXT ou DOCX).

    Retorna (texto_completo, lista_de_sentenças).
    """
    filename = uploaded_file.name.lower()
    raw_bytes = uploaded_file.read()

    if filename.endswith(".pdf"):
        doc = fitz.open(stream=raw_bytes, filetype="pdf")
        text = ""
        for page in doc:
            for block in page.get_text("blocks"):
                cleaned = clear_texts(block[4])
                if cleaned:
                    text += cleaned + " "
        doc.close()
    elif filename.endswith(".txt"):
        text = raw_bytes.decode("utf-8", errors="ignore")
    elif filename.endswith(".docx"):
        text = _extract_docx(raw_bytes)
    else:
        raise ValueError(
            f"Tipo de arquivo não suportado: '{uploaded_file.name}'.\n"
            "Formatos aceitos: PDF, TXT, DOCX."
        )

    text = text.strip()
    sentences = [s.strip() for s in sent_tokenize(text) if len(s.strip()) > 30]
    return text, sentences












